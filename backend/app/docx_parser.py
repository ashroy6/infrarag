from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


DOCX_EXTENSIONS = {".docx"}


def supports(path: str | Path) -> bool:
    return Path(path).suffix.lower() in DOCX_EXTENSIONS


def _extract_paragraphs(document: Document) -> list[str]:
    paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = " ".join((paragraph.text or "").split())
        if text:
            paragraphs.append(text)

    return paragraphs


def _extract_tables(document: Document) -> list[str]:
    table_lines: list[str] = []

    for table_index, table in enumerate(document.tables, start=1):
        table_lines.append(f"[Table {table_index}]")

        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join((cell.text or "").split())
                cells.append(cell_text)

            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                table_lines.append(row_text)

    return table_lines


def parse_docx_file(path: str | Path) -> dict[str, Any]:
    p = Path(path)
    document = Document(str(p))

    paragraphs = _extract_paragraphs(document)
    tables = _extract_tables(document)

    sections: list[str] = []
    if paragraphs:
        sections.append("\n".join(paragraphs))
    if tables:
        sections.append("\n".join(tables))

    text = "\n\n".join(section for section in sections if section.strip()).strip()

    return {
        "path": str(p),
        "parser_type": "docx",
        "file_type": ".docx",
        "text": text,
        "metadata": {
            "filename": p.name,
            "parent": str(p.parent),
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
        },
    }
